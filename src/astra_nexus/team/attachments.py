from __future__ import annotations

import shutil
from dataclasses import dataclass, field
from enum import StrEnum
from pathlib import Path
from typing import Any


class TeamAttachmentType(StrEnum):
    TEXT = "text"
    MARKDOWN = "markdown"
    PDF = "pdf"
    DOCX = "docx"
    BINARY = "binary"
    UNKNOWN = "unknown"


class TeamAttachmentExtractionStatus(StrEnum):
    NOT_NEEDED = "not_needed"
    EXTRACTED = "extracted"
    METADATA_ONLY = "metadata_only"
    FAILED = "failed"
    TRUNCATED = "truncated"
    PENDING = "not_needed"
    UNSUPPORTED = "metadata_only"
    ERROR = "failed"


class TeamAttachmentValidationError(ValueError):
    pass


@dataclass
class TeamInputAttachment:
    original_filename: str
    stored_filename: str
    content_type: str | None = None
    size_bytes: int = 0
    source: str = "local"
    local_path: Path | None = None
    attachment_type: TeamAttachmentType = TeamAttachmentType.UNKNOWN
    extracted_text: str | None = None
    extraction_status: TeamAttachmentExtractionStatus = TeamAttachmentExtractionStatus.PENDING
    extraction_error: str | None = None
    metadata: dict[str, Any] = field(default_factory=dict)
    extension: str | None = None
    mime_type: str | None = None
    extracted_chars: int = 0
    prompt_chars: int = 0
    pages_count: int | None = None
    paragraphs_count: int | None = None
    truncated: bool = False

    @property
    def prompt_text(self) -> str:
        if not self.extracted_text:
            return ""
        if self.prompt_chars <= 0:
            return ""
        return self.extracted_text[: self.prompt_chars].rstrip()


@dataclass(frozen=True)
class TeamAttachmentManifest:
    attachments: tuple[TeamInputAttachment, ...] = ()

    @property
    def count(self) -> int:
        return len(self.attachments)


class TeamAttachmentProcessor:
    def __init__(
        self,
        *,
        max_files: int = 5,
        max_bytes: int = 10 * 1024 * 1024,
        text_max_chars: int | None = None,
        max_extracted_chars: int = 50000,
        max_prompt_chars: int = 20000,
        pdf_max_pages: int = 30,
        docx_include_tables: bool = True,
    ) -> None:
        self.max_files = max_files
        self.max_bytes = max_bytes
        if text_max_chars is not None:
            max_extracted_chars = text_max_chars
            max_prompt_chars = text_max_chars
        self.max_extracted_chars = max_extracted_chars
        self.max_prompt_chars = max_prompt_chars
        self.pdf_max_pages = pdf_max_pages
        self.docx_include_tables = docx_include_tables

    def prepare_paths(
        self,
        paths: list[Path] | tuple[Path, ...],
        *,
        source: str,
    ) -> tuple[TeamInputAttachment, ...]:
        self._validate_count(len(paths))
        attachments = []
        for path in paths:
            path = Path(path)
            size_bytes = path.stat().st_size
            self._validate_size(path.name, size_bytes)
            attachments.append(
                TeamInputAttachment(
                    original_filename=path.name,
                    stored_filename=_safe_filename(path.name),
                    content_type=_content_type(path),
                    size_bytes=size_bytes,
                    source=source,
                    local_path=path,
                    attachment_type=_attachment_type(path),
                    extension=_extension(path),
                    mime_type=_content_type(path),
                )
            )
        return self.process(attachments)

    def process(
        self,
        attachments: list[TeamInputAttachment] | tuple[TeamInputAttachment, ...],
    ) -> tuple[TeamInputAttachment, ...]:
        self._validate_count(len(attachments))
        for attachment in attachments:
            self._validate_size(attachment.original_filename, attachment.size_bytes)
            self._extract(attachment)
        return tuple(attachments)

    def _extract(self, attachment: TeamInputAttachment) -> None:
        if attachment.attachment_type == TeamAttachmentType.UNKNOWN and attachment.local_path:
            attachment.attachment_type = _attachment_type(attachment.local_path)
        self._fill_basic_metadata(attachment)

        if attachment.attachment_type in {TeamAttachmentType.BINARY, TeamAttachmentType.UNKNOWN}:
            attachment.extraction_status = TeamAttachmentExtractionStatus.METADATA_ONLY
            attachment.extracted_text = None
            self._sync_metadata(attachment)
            return

        if attachment.local_path is None:
            self._fail(attachment, "local_path is missing")
            return

        if attachment.attachment_type == TeamAttachmentType.TEXT:
            self._extract_text_file(attachment)
            return
        if attachment.attachment_type == TeamAttachmentType.MARKDOWN:
            self._extract_text_file(attachment)
            return
        if attachment.attachment_type == TeamAttachmentType.DOCX:
            self._extract_docx(attachment)
            return
        if attachment.attachment_type == TeamAttachmentType.PDF:
            self._extract_pdf(attachment)
            return

        attachment.extraction_status = TeamAttachmentExtractionStatus.METADATA_ONLY
        attachment.extracted_text = None
        self._sync_metadata(attachment)

    def _extract_text_file(self, attachment: TeamInputAttachment) -> None:
        try:
            text = attachment.local_path.read_text(encoding="utf-8")
        except Exception as exc:
            self._fail(attachment, str(exc))
            return
        self._set_extracted_text(attachment, text)

    def _extract_docx(self, attachment: TeamInputAttachment) -> None:
        try:
            from docx import Document

            document = Document(str(attachment.local_path))
            paragraphs = [
                paragraph.text.strip()
                for paragraph in document.paragraphs
                if paragraph.text.strip()
            ]
            sections = list(paragraphs)
            if self.docx_include_tables:
                for table in document.tables:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if cells:
                            sections.append(" | ".join(cells))
            attachment.paragraphs_count = len(paragraphs)
            attachment.metadata["tables_count"] = len(document.tables)
        except Exception as exc:
            self._fail(attachment, str(exc))
            return
        self._set_extracted_text(attachment, "\n\n".join(sections))

    def _extract_pdf(self, attachment: TeamInputAttachment) -> None:
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(attachment.local_path))
            attachment.pages_count = len(reader.pages)
            pages = reader.pages[: max(0, self.pdf_max_pages)]
            sections: list[str] = []
            for index, page in enumerate(pages, start=1):
                sections.append(f"Страница {index}")
                try:
                    page_text = page.extract_text() or ""
                except Exception as exc:
                    page_text = f"[Текст страницы не извлечён: {exc}]"
                sections.append(page_text.strip() or "[Текст страницы не извлечён]")
            if attachment.pages_count > len(pages):
                sections.append(
                    f"[PDF обработан до {len(pages)} страниц из {attachment.pages_count}.]"
                )
        except Exception as exc:
            self._fail(attachment, str(exc))
            return
        self._set_extracted_text(
            attachment,
            "\n\n".join(sections),
            source_truncated=attachment.pages_count is not None
            and attachment.pages_count > max(0, self.pdf_max_pages),
        )

    def _set_extracted_text(
        self,
        attachment: TeamInputAttachment,
        text: str,
        *,
        source_truncated: bool = False,
    ) -> None:
        extracted_text = text
        extracted_limit = self.max_extracted_chars
        extraction_truncated = extracted_limit > 0 and len(extracted_text) > extracted_limit
        if extraction_truncated:
            extracted_text = extracted_text[:extracted_limit].rstrip()

        prompt_limit = self.max_prompt_chars
        if prompt_limit > 0:
            prompt_chars = min(len(extracted_text), prompt_limit)
            prompt_truncated = len(extracted_text) > prompt_chars
        else:
            prompt_chars = len(extracted_text)
            prompt_truncated = False

        attachment.extracted_text = extracted_text
        attachment.extracted_chars = len(extracted_text)
        attachment.prompt_chars = prompt_chars
        attachment.truncated = bool(source_truncated or extraction_truncated or prompt_truncated)
        attachment.extraction_status = (
            TeamAttachmentExtractionStatus.TRUNCATED
            if attachment.truncated
            else TeamAttachmentExtractionStatus.EXTRACTED
        )
        attachment.extraction_error = None
        self._sync_metadata(attachment)

    def _fill_basic_metadata(self, attachment: TeamInputAttachment) -> None:
        path = attachment.local_path
        if attachment.extension is None:
            attachment.extension = _extension(path or Path(attachment.original_filename))
        if attachment.mime_type is None:
            attachment.mime_type = attachment.content_type or _content_type(
                path or Path(attachment.original_filename)
            )
        if attachment.content_type is None:
            attachment.content_type = attachment.mime_type

    def _fail(self, attachment: TeamInputAttachment, error: str) -> None:
        attachment.extraction_status = TeamAttachmentExtractionStatus.FAILED
        attachment.extraction_error = error
        attachment.extracted_text = None
        attachment.extracted_chars = 0
        attachment.prompt_chars = 0
        attachment.truncated = False
        self._sync_metadata(attachment)

    def _sync_metadata(self, attachment: TeamInputAttachment) -> None:
        attachment.metadata.update(
            {
                "extension": attachment.extension,
                "mime_type": attachment.mime_type or attachment.content_type,
                "original_size": attachment.size_bytes,
                "extracted_chars": attachment.extracted_chars,
                "prompt_chars": attachment.prompt_chars,
                "pages_count": attachment.pages_count,
                "paragraphs_count": attachment.paragraphs_count,
                "truncated": attachment.truncated,
            }
        )

    def _validate_count(self, count: int) -> None:
        if count > self.max_files:
            raise TeamAttachmentValidationError(
                f"too many attachments: {count}; max allowed is {self.max_files}"
            )

    def _validate_size(self, filename: str, size_bytes: int) -> None:
        if size_bytes > self.max_bytes:
            raise TeamAttachmentValidationError(
                f"attachment {filename} is too large: {size_bytes}; max allowed is {self.max_bytes}"
            )


def save_attachments_to_workspace(
    attachments: list[TeamInputAttachment],
    *,
    run_path: Path,
) -> None:
    input_files_path = run_path / "input_files"
    input_files_path.mkdir(parents=True, exist_ok=True)
    used_names: set[str] = set()
    for index, attachment in enumerate(attachments, start=1):
        stored_filename = _unique_filename(
            attachment.stored_filename or _safe_filename(attachment.original_filename),
            used_names,
            index=index,
        )
        attachment.stored_filename = stored_filename
        destination = input_files_path / stored_filename
        if attachment.local_path is not None and attachment.local_path.exists():
            source = attachment.local_path
            if source.resolve() != destination.resolve():
                shutil.copy2(source, destination)
            attachment.local_path = destination


def attachment_payload(attachment: TeamInputAttachment) -> dict[str, Any]:
    return {
        "original_filename": attachment.original_filename,
        "stored_filename": attachment.stored_filename,
        "content_type": attachment.content_type,
        "mime_type": attachment.mime_type or attachment.content_type,
        "extension": attachment.extension,
        "size_bytes": attachment.size_bytes,
        "original_size": attachment.size_bytes,
        "source": attachment.source,
        "local_path": str(attachment.local_path) if attachment.local_path else None,
        "attachment_type": attachment.attachment_type.value,
        "extracted_text": attachment.extracted_text,
        "extraction_status": attachment.extraction_status.value,
        "extraction_error": attachment.extraction_error,
        "extracted_chars": attachment.extracted_chars,
        "prompt_chars": attachment.prompt_chars,
        "pages_count": attachment.pages_count,
        "paragraphs_count": attachment.paragraphs_count,
        "truncated": attachment.truncated,
        "metadata": attachment.metadata,
    }


def attachment_from_payload(payload: dict[str, Any]) -> TeamInputAttachment:
    extracted_text = payload.get("extracted_text")
    extracted_chars = payload.get("extracted_chars")
    if extracted_chars is None:
        extracted_chars = len(extracted_text or "")
    prompt_chars = payload.get("prompt_chars")
    if prompt_chars is None or (prompt_chars == 0 and extracted_text):
        prompt_chars = extracted_chars
    filename_for_metadata = payload.get("stored_filename") or payload.get("original_filename", "")
    return TeamInputAttachment(
        original_filename=payload["original_filename"],
        stored_filename=payload["stored_filename"],
        content_type=payload.get("content_type"),
        size_bytes=payload.get("size_bytes", 0),
        source=payload.get("source", "workspace"),
        local_path=Path(payload["local_path"]) if payload.get("local_path") else None,
        attachment_type=TeamAttachmentType(payload.get("attachment_type", "unknown")),
        extracted_text=extracted_text,
        extraction_status=_extraction_status(payload.get("extraction_status", "metadata_only")),
        extraction_error=payload.get("extraction_error"),
        metadata=payload.get("metadata", {}),
        extension=payload.get("extension") or _extension(Path(filename_for_metadata)),
        mime_type=payload.get("mime_type") or payload.get("content_type"),
        extracted_chars=extracted_chars,
        prompt_chars=prompt_chars,
        pages_count=payload.get("pages_count"),
        paragraphs_count=payload.get("paragraphs_count"),
        truncated=payload.get("truncated", False),
    )


def attachments_markdown(attachments: list[TeamInputAttachment]) -> str:
    sections = ["# Attachments", ""]
    if not attachments:
        sections.extend(["No attachments.", ""])
        return "\n".join(sections)

    for index, attachment in enumerate(attachments, start=1):
        sections.extend(
            [
                f"## {index}. {attachment.original_filename}",
                "",
                f"- Stored filename: `{attachment.stored_filename}`",
                f"- Content type: `{attachment.content_type or 'unknown'}`",
                f"- MIME type: `{attachment.mime_type or attachment.content_type or 'unknown'}`",
                f"- Extension: `{attachment.extension or ''}`",
                f"- Size: `{attachment.size_bytes}` bytes",
                f"- Source: `{attachment.source}`",
                f"- Local path: `{attachment.local_path or ''}`",
                f"- Extraction status: `{attachment.extraction_status.value}`",
                f"- Extracted chars: `{attachment.extracted_chars}`",
                f"- Prompt chars: `{attachment.prompt_chars}`",
                f"- Truncated: `{attachment.truncated}`",
            ]
        )
        if attachment.pages_count is not None:
            sections.append(f"- Pages count: `{attachment.pages_count}`")
        if attachment.paragraphs_count is not None:
            sections.append(f"- Paragraphs count: `{attachment.paragraphs_count}`")
        if attachment.extraction_error:
            sections.append(f"- Extraction error: `{attachment.extraction_error}`")
        if attachment.extracted_text:
            sections.extend(
                ["", "### Extracted text", "", "```text", attachment.extracted_text, "```"]
            )
        sections.append("")
    return "\n".join(sections).rstrip() + "\n"


def _attachment_type(path: Path) -> TeamAttachmentType:
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return TeamAttachmentType.TEXT
    if suffix == ".md":
        return TeamAttachmentType.MARKDOWN
    if suffix == ".pdf":
        return TeamAttachmentType.PDF
    if suffix == ".docx":
        return TeamAttachmentType.DOCX
    if suffix:
        return TeamAttachmentType.BINARY
    return TeamAttachmentType.UNKNOWN


def _content_type(path: Path) -> str | None:
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return "text/plain"
    if suffix == ".md":
        return "text/markdown"
    if suffix == ".pdf":
        return "application/pdf"
    if suffix == ".docx":
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return None


def _extension(path: Path) -> str:
    return path.suffix.lower()


def _extraction_status(value: Any) -> TeamAttachmentExtractionStatus:
    legacy = {
        "pending": TeamAttachmentExtractionStatus.NOT_NEEDED,
        "unsupported": TeamAttachmentExtractionStatus.METADATA_ONLY,
        "error": TeamAttachmentExtractionStatus.FAILED,
    }
    text = str(value or "metadata_only")
    return legacy.get(text, TeamAttachmentExtractionStatus(text))


def _safe_filename(filename: str) -> str:
    safe = "".join(char if char.isalnum() or char in "._-" else "_" for char in filename)
    safe = safe.strip("._")
    return safe or "attachment"


def _unique_filename(filename: str, used_names: set[str], *, index: int) -> str:
    candidate = filename
    if candidate not in used_names:
        used_names.add(candidate)
        return candidate
    path = Path(filename)
    suffix = path.suffix
    stem = path.stem or "attachment"
    while candidate in used_names:
        candidate = f"{stem}_{index}{suffix}"
        index += 1
    used_names.add(candidate)
    return candidate
